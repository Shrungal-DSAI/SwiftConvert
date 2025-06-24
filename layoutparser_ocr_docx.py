import layoutparser as lp
import pytesseract
import cv2
from docx import Document
import os

# Set tesseract path if not in system PATH
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

def ocr_layout_to_docx(image_path, output_docx_path):
    image = cv2.imread(image_path)
    if image is None:
        raise FileNotFoundError(f"Image not found: {image_path}")

    # CPU-friendly model from LayoutParser (detectron2)
    model = lp.Detectron2LayoutModel(
        config_path="lp://PubLayNet/faster_rcnn_R_50_FPN_3x/config",
        label_map={0: "Text", 1: "Title", 2: "List", 3: "Table", 4: "Figure"},
        extra_config=["MODEL.ROI_HEADS.SCORE_THRESH_TEST", 0.5]
    )

    layout = model.detect(image)
    layout = lp.Layout([b for b in layout if b.type != 'Figure']).sort(key=lambda b: b.block.y_1)

    doc = Document()
    doc.add_heading("OCR Extracted Content", level=1)

    for block in layout:
        segment_image = block.crop_image(image)
        text = pytesseract.image_to_string(segment_image)
        if text.strip():
            doc.add_paragraph(f"[{block.type}] {text.strip()}")

    doc.save(output_docx_path)
    print(f"DOCX saved: {output_docx_path}")

if __name__ == "__main__":
    image_file = "example_invoice.png"  # make sure this file exists
    output_file = "output_result.docx"
    ocr_layout_to_docx(image_file, output_file)
