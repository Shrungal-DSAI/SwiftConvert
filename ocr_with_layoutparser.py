import layoutparser as lp
import cv2
import pytesseract
from pdf2image import convert_from_path
from docx import Document
import os

def convert_pdf_to_docx_with_ocr(pdf_path, output_docx_path):
    # Convert PDF to images
    images = convert_from_path(pdf_path, dpi=300)
    document = Document()

    model = lp.Detectron2LayoutModel(
        config_path='lp://PubLayNet/faster_rcnn_R_50_FPN_3x/config',
        model_path=None,
        label_map={0: "Text", 1: "Title", 2: "List", 3: "Table", 4: "Figure"},
        extra_config=["MODEL.ROI_HEADS.SCORE_THRESH_TEST", 0.8],
    )

    for i, image in enumerate(images):
        image_np = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
        layout = model.detect(image_np)

        for block in layout:
            segment_image = block.crop_image(image_np)
            text = pytesseract.image_to_string(segment_image, lang="eng")

            if block.type == "Title":
                document.add_heading(text.strip(), level=1)
            elif block.type == "Table":
                document.add_paragraph("[TABLE DETECTED]")
                document.add_paragraph(text.strip())
            elif block.type == "List":
                document.add_paragraph(f"• {text.strip()}")
            else:
                document.add_paragraph(text.strip())

        if i < len(images) - 1:
            document.add_page_break()

    document.save(output_docx_path)
