from pdf2image import convert_from_path
import pytesseract
from pytesseract import Output
from docx import Document
from PIL import Image
import os
from pathlib import Path
import logging

logging.basicConfig(level=logging.INFO)

def convert_pdf_to_docx_with_ocr(pdf_path, docx_output_path):
    try:
        logging.info(f"📄 Starting OCR conversion for: {pdf_path}")

        # Ensure output directory exists
        output_dir = Path(docx_output_path).parent
        output_dir.mkdir(parents=True, exist_ok=True)

        # Convert PDF to image pages
        pages = convert_from_path(pdf_path, dpi=300)
        doc = Document()

        for i, page in enumerate(pages):
            doc.add_paragraph(f"📄 Page {i + 1}")
            data = pytesseract.image_to_data(page, output_type=Output.DICT)

            block_text = ""
            prev_block_num = None

            for j, word in enumerate(data["text"]):
                try:
                    if int(data["conf"][j]) > 30 and word.strip():
                        block_num = data["block_num"][j]

                        if block_num != prev_block_num and block_text:
                            doc.add_paragraph(block_text.strip())
                            block_text = ""

                        block_text += word + " "
                        prev_block_num = block_num
                except (ValueError, IndexError):
                    continue  # skip malformed data

            if block_text:
                doc.add_paragraph(block_text.strip())

            doc.add_page_break()

        doc.save(docx_output_path)
        logging.info(f"✅ OCR to DOCX complete: {docx_output_path}")

    except Exception as e:
        logging.error(f"❌ OCR Conversion failed: {str(e)}")
        raise RuntimeError(f"OCR Conversion failed: {str(e)}")