import fitz  # PyMuPDF
import pdfplumber
from docx import Document
import os
import logging

def convert_pdf_to_docx_without_ocr(pdf_path, docx_output_path):
    logging.info(f"📄 Starting non-OCR conversion for: {pdf_path}")
    doc = Document()

    # Step 1: Extract Text with PyMuPDF
    with fitz.open(pdf_path) as pdf:
        for i, page in enumerate(pdf):
            text = page.get_text("text")
            doc.add_paragraph(f"📄 Page {i + 1}")
            doc.add_paragraph(text)
            doc.add_page_break()

    # Step 2: Extract Tables with pdfplumber
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            for table in tables:
                doc.add_paragraph(f"📊 Table on Page {i + 1}")
                for row in table:
                    doc.add_paragraph(" | ".join(cell if cell else "" for cell in row))

    doc.save(docx_output_path)
    logging.info(f"✅ Non-OCR PDF to DOCX complete: {docx_output_path}")
