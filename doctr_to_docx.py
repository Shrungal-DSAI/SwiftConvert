from doctr.io import DocumentFile
from doctr.models import ocr_predictor
from docx import Document
import os

# Load the document
pdf_path = r"C:\Users\Admin\Downloads\Antragsformular_0af903fb_4572_441b_8fcd_e4f19b222ae4.pdf"
doc = DocumentFile.from_pdf(pdf_path)

# Load OCR model
model = ocr_predictor(pretrained=True)

# Perform OCR
result = model(doc)

# Export to DOCX
output_doc = Document()

for i, page in enumerate(result.pages):
    output_doc.add_paragraph(f"--- Page {i + 1} ---")
    for block in page.blocks:
        for line in block.lines:
            line_text = " ".join([word.value for word in line.words])
            output_doc.add_paragraph(line_text)
    output_doc.add_page_break()

output_path = "ocr_output.docx"
output_doc.save(output_path)
print(f"✅ Exported to {output_path}")
